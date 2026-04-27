#!/usr/bin/env python3
"""
Scraper completo da Media Library do Portal da Transparência do CAU/PR.

Usa a WordPress REST API /wp/v2/media para descobrir TODOS os arquivos
enviados (PDFs, DOC/DOCX, XLS, imagens, ZIPs) e os insere no banco,
extraindo texto de PDFs e DOCX.

Tipos tratados:
  PDF          → pdfplumber (texto nativo) ou marca como escaneado
  DOC / DOCX   → python-docx
  XLS / XLSX   → registra sem extração (formato_xls)
  ZIP          → registra sem extração (formato_zip)
  Imagem (JPG/PNG/TIFF) → marca como escaneado (futuro OCR)
  HTML         → registra sem extração

Uso (cd backend/):
    python scripts/scrape_media_library.py --dry-run          # lista sem baixar
    python scripts/scrape_media_library.py --dry-run --tipo pdf
    python scripts/scrape_media_library.py                    # baixa tudo
    python scripts/scrape_media_library.py --tipo pdf         # só PDFs
    python scripts/scrape_media_library.py --tipo docx        # só DOC/DOCX
    python scripts/scrape_media_library.py --tipo imagem      # só imagens
    python scripts/scrape_media_library.py --ano 2021         # só uploads de 2021
    python scripts/scrape_media_library.py --limit 50         # primeiros 50 arquivos
"""
import argparse
import asyncio
import html as hlib
import io
import os
import re
import sys
import uuid
from pathlib import Path

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

from dotenv import load_dotenv
load_dotenv(ROOT / ".env")

import asyncpg
import httpx

try:
    import pdfplumber
    import ftfy
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ── Config ─────────────────────────────────────────────────────────────────────
TENANT_ID  = "f32ed0e7-c95d-4dec-a332-fa2cf6f20eb4"
WP_BASE    = "https://transparencia.caupr.gov.br/wp-json/wp/v2"
RATE_LIMIT = 1.5   # segundos entre downloads
MAX_BYTES  = 100 * 1024 * 1024  # 100 MB

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Referer":         "https://transparencia.caupr.gov.br/",
    "Accept":          "*/*",
    "Accept-Language": "pt-BR,pt;q=0.9",
}

DATABASE_URL = os.environ.get("DATABASE_URL", "")
if not DATABASE_URL:
    sys.exit("ERROR: DATABASE_URL não encontrado no backend/.env")
ASYNCPG_URL = DATABASE_URL.replace("postgresql+asyncpg://", "postgresql://")

# Mapeamento mime_type → categoria
MIME_MAP = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
    "application/vnd.ms-excel": "xls",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/zip": "zip",
    "application/x-zip-compressed": "zip",
    "text/html": "html",
    "image/jpeg": "imagem",
    "image/png":  "imagem",
    "image/tiff": "imagem",
    "image/gif":  "imagem",
    "image/webp": "imagem",
}

TIPO_FILTRO_MAP = {
    "pdf":    ["application/pdf"],
    "docx":   ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"],
    "doc":    ["application/msword"],
    "xls":    ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"],
    "imagem": ["image/jpeg", "image/png", "image/tiff", "image/gif", "image/webp"],
    "zip":    ["application/zip", "application/x-zip-compressed"],
}


# ── Helpers ────────────────────────────────────────────────────────────────────

def _filename(url: str) -> str:
    return url.split('?')[0].rstrip('/').split('/')[-1]


def _numero(url: str) -> str:
    path = url.split('?')[0].rstrip('/')
    fn   = path.split('/')[-1]
    name = re.sub(r'\.(pdf|xls|xlsx|doc|docx|zip|jpg|jpeg|png|tiff|gif|html?)$', '', fn, flags=re.I)
    name = re.sub(r'[%+\s]', '-', name)
    name = re.sub(r'-{2,}', '-', name).strip('-')
    return name[:100] or fn[:100]


def _extrair_pdf(pdf_bytes: bytes) -> tuple[str, int]:
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        n_pages = len(pdf.pages)
        parts   = [p.extract_text() or "" for p in pdf.pages]
    texto = ftfy.fix_text("\n\n".join(parts))
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip(), n_pages


def _extrair_docx(docx_bytes: bytes) -> str:
    doc    = DocxDocument(io.BytesIO(docx_bytes))
    partes = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            partes.append(t)
    for table in doc.tables:
        for row in table.rows:
            celulas = [c.text.strip() for c in row.cells if c.text.strip()]
            if celulas:
                partes.append(" | ".join(celulas))
    return "\n".join(partes)


# ── Fetch media library ────────────────────────────────────────────────────────

async def _fetch_media_page(
    http:      httpx.AsyncClient,
    mime:      str | None,
    page:      int,
    ano_after: str | None,
    ano_before: str | None,
) -> list[dict]:
    params = f"per_page=100&page={page}&_fields=id,date,source_url,title,mime_type"
    if mime:
        params += f"&mime_type={mime}"
    if ano_after:
        params += f"&after={ano_after}"
    if ano_before:
        params += f"&before={ano_before}"

    url = f"{WP_BASE}/media?{params}"
    try:
        r = await http.get(url, timeout=30)
        if r.status_code in (400, 404):
            return []
        r.raise_for_status()
        return r.json()
    except Exception as e:
        print(f"  ⚠  API erro (p{page}): {e}", file=sys.stderr)
        return []


async def descobrir_todos(
    http:       httpx.AsyncClient,
    mimes:      list[str] | None,
    ano_ini:    int | None,
    ano_fim:    int | None,
) -> list[dict]:
    """Retorna lista de todos os arquivos na media library."""
    after  = f"{ano_ini}-01-01T00:00:00" if ano_ini else None
    before = f"{ano_fim}-12-31T23:59:59" if ano_fim else None

    todos: list[dict] = []
    vistos: set[str]  = set()

    mime_list = mimes if mimes else [None]  # None = sem filtro (todos)

    for mime in mime_list:
        page = 1
        while True:
            items = await _fetch_media_page(http, mime, page, after, before)
            if not items:
                break
            for item in items:
                url = item.get("source_url", "")
                if url and url not in vistos:
                    vistos.add(url)
                    todos.append(item)
            page += 1
            await asyncio.sleep(0.3)

    return todos


# ── Download + extração ────────────────────────────────────────────────────────

async def _processar_item(
    conn:     asyncpg.Connection,
    http:     httpx.AsyncClient,
    item:     dict,
    dry_run:  bool,
    idx:      int,
    total:    int,
) -> str:
    """Retorna código: 'ok' | 'ja_existe' | 'escaneado' | 'nao_extracao' | 'erro'"""
    url      = item["source_url"]
    mime     = item.get("mime_type", "")
    data_wp  = item.get("date", "")[:10]
    titulo   = hlib.unescape(item.get("title", {}).get("rendered", ""))
    numero   = _numero(url)
    fn       = _filename(url)
    categoria = MIME_MAP.get(mime, "outro")

    prefix = f"[{idx:4d}/{total}] {fn[:50]}"

    if dry_run:
        print(f"  [{categoria:6s}] {data_wp}  {fn[:70]}")
        return "dry_run"

    # INSERT idempotente em atos
    ato_row = await conn.fetchrow(
        """
        INSERT INTO atos
            (id, tenant_id, numero, tipo, subtipo,
             url_original, url_pdf, pdf_baixado, processado, ementa)
        VALUES ($1, $2, $3, 'media_library', $4, $5, $6, false, false, $7)
        ON CONFLICT (tenant_id, numero, tipo) DO NOTHING
        RETURNING id
        """,
        uuid.uuid4(), TENANT_ID, numero, categoria,
        f"https://transparencia.caupr.gov.br/",
        url,
        titulo[:500] if titulo else fn,
    )

    if ato_row is None:
        ato_row = await conn.fetchrow(
            "SELECT id FROM atos WHERE tenant_id=$1 AND numero=$2 AND tipo='media_library'",
            TENANT_ID, numero,
        )
        tem_conteudo = await conn.fetchval(
            "SELECT 1 FROM conteudo_ato WHERE ato_id=$1", ato_row["id"]
        )
        if tem_conteudo:
            print(f"  ─  {prefix}: já no banco com conteúdo")
            return "ja_existe"

    ato_id = ato_row["id"]

    # Imagem → marca escaneado sem download
    if categoria == "imagem":
        await conn.execute(
            "UPDATE atos SET erro_download='escaneado_imagem' WHERE id=$1", ato_id
        )
        print(f"  🖼  {prefix}: imagem (futuro OCR)")
        return "escaneado"

    # Não-extração (XLS, ZIP, HTML) → registra sem download
    if categoria in ("xls", "xlsx", "zip", "html", "outro"):
        await conn.execute(
            "UPDATE atos SET erro_download=$1 WHERE id=$2", f"formato_{categoria}", ato_id
        )
        print(f"  ─  {prefix}: [{categoria}] registrado sem extração")
        return "nao_extracao"

    # PDF / DOCX / DOC → baixa e extrai
    try:
        resp = await http.get(url, timeout=90)
        if resp.status_code == 403:
            await conn.execute(
                "UPDATE atos SET erro_download='403_forbidden' WHERE id=$1", ato_id
            )
            print(f"  ⛔ {prefix}: 403 Forbidden")
            return "erro"
        resp.raise_for_status()
        raw = resp.content

        if len(raw) > MAX_BYTES:
            raise ValueError(f"arquivo muito grande: {len(raw):,} bytes")

        texto = ""
        n_pages = 0
        metodo  = "desconhecido"

        if categoria == "pdf":
            if not HAS_PDF:
                raise ImportError("pdfplumber não instalado")
            texto, n_pages = _extrair_pdf(raw)
            metodo = "pdfplumber"

        elif categoria in ("docx", "doc"):
            if not HAS_DOCX:
                raise ImportError("python-docx não instalado")
            texto = _extrair_docx(raw)
            metodo = "python-docx"
            n_pages = max(1, len(texto) // 3000)

        qualidade = "boa" if len(texto) > 100 else "ruim"
        tokens    = len(texto) // 4

        await conn.execute(
            """
            INSERT INTO conteudo_ato
                (ato_id, texto_completo, metodo_extracao, qualidade, tokens_estimados, criado_em)
            VALUES ($1, $2, $3, $4, $5, NOW())
            ON CONFLICT (ato_id) DO NOTHING
            """,
            ato_id, texto, metodo, qualidade, tokens,
        )
        await conn.execute(
            """
            UPDATE atos SET
                pdf_baixado       = true,
                pdf_paginas       = $1,
                pdf_tamanho_bytes = $2
            WHERE id = $3
            """,
            n_pages, len(raw), ato_id,
        )

        if qualidade == "boa":
            print(f"  ✓  {prefix}: {n_pages}p  {tokens:,} tokens [{metodo}]")
            return "ok"
        else:
            print(f"  ⚠  {prefix}: {n_pages}p — escaneado (sem texto)")
            return "escaneado"

    except Exception as exc:
        await conn.execute(
            "UPDATE atos SET erro_download=$1 WHERE id=$2", str(exc)[:200], ato_id
        )
        print(f"  ✗  {prefix}: {exc}")
        return "erro"


# ── Main ────────────────────────────────────────────────────────────────────────

async def main(args: argparse.Namespace) -> None:
    print(f"\n{'='*70}")
    print("Scraper — Media Library Completa — Portal da Transparência CAU/PR")
    print(f"{'='*70}\n")

    # Resolve filtros de tipo
    mimes: list[str] | None = None
    if args.tipo:
        tipos = [t.strip() for t in args.tipo.split(",")]
        mimes = []
        for t in tipos:
            mimes.extend(TIPO_FILTRO_MAP.get(t, []))
        if not mimes:
            sys.exit(f"Tipo(s) inválido(s). Válidos: {', '.join(TIPO_FILTRO_MAP)}")

    ano_ini = ano_fim = None
    if args.ano:
        partes  = args.ano.split("-")
        ano_ini = int(partes[0])
        ano_fim = int(partes[-1])

    modo = "DRY RUN" if args.dry_run else "DOWNLOAD"
    print(f"Modo:    {modo}")
    print(f"Tipos:   {args.tipo or 'todos'}")
    print(f"Ano:     {args.ano or 'todos'}")
    print(f"Limit:   {args.limit or 'sem limite'}\n")

    if not args.dry_run and not HAS_PDF:
        print("⚠  pdfplumber não instalado — PDFs serão registrados sem extração de texto")
    if not args.dry_run and not HAS_DOCX:
        print("⚠  python-docx não instalado — DOC/DOCX serão registrados sem extração de texto")

    async with httpx.AsyncClient(headers=HEADERS, follow_redirects=True) as http:
        print("Consultando media library...", flush=True)
        todos = await descobrir_todos(http, mimes, ano_ini, ano_fim)

        if args.limit:
            todos = todos[:args.limit]

        print(f"Total encontrado: {len(todos)} arquivos\n")

        if args.dry_run:
            for item in todos:
                mime = item.get("mime_type", "")
                cat  = MIME_MAP.get(mime, "outro")
                data = item.get("date", "")[:10]
                fn   = _filename(item.get("source_url", ""))
                print(f"  [{cat:6s}] {data}  {fn}")
            print(f"\nTotal: {len(todos)} arquivos (dry run — nenhum download feito)")
            return

        # Contadores
        totais = {"ok": 0, "ja_existe": 0, "escaneado": 0, "nao_extracao": 0, "erro": 0}

        conn = await asyncpg.connect(ASYNCPG_URL, statement_cache_size=0)
        try:
            for i, item in enumerate(todos, 1):
                resultado = await _processar_item(conn, http, item, False, i, len(todos))
                totais[resultado] = totais.get(resultado, 0) + 1
                if i < len(todos):
                    await asyncio.sleep(RATE_LIMIT)
        finally:
            await conn.close()

    print(f"\n{'='*70}")
    print("TOTAIS FINAIS")
    print(f"{'='*70}")
    print(f"  ✓  Com texto extraído:        {totais['ok']}")
    print(f"  ─  Já existiam no banco:      {totais['ja_existe']}")
    print(f"  ─  Formato sem extração:      {totais['nao_extracao']}")
    print(f"  ⚠  Escaneados (futuro OCR):   {totais['escaneado']}")
    print(f"  ✗  Erros:                     {totais['erro']}")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Scraper completo da media library do Portal da Transparência CAU/PR"
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="Lista arquivos sem baixar nem inserir no banco"
    )
    parser.add_argument(
        "--tipo", type=str, default="",
        help="Filtra por tipo: pdf, docx, doc, xls, imagem, zip (separados por vírgula)"
    )
    parser.add_argument(
        "--ano", type=str, default="",
        help="Filtra por ano ou intervalo: 2021 ou 2019-2023"
    )
    parser.add_argument(
        "--limit", type=int, default=0,
        help="Limita o número de arquivos processados (útil para testes)"
    )
    args = parser.parse_args()

    asyncio.run(main(args))
