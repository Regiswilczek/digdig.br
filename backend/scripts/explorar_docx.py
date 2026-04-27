#!/usr/bin/env python3
"""
Explorador de arquivos .docx no Portal da Transparência do CAU/PR.

Modos de descoberta:
  - Seções linkadas: varre as 21 seções mapeadas (e sub-páginas com --deep)
  - Media library (--media): consulta WP REST API /wp/v2/media para encontrar
    TODOS os .docx enviados, mesmo os não linkados em nenhuma página (órfãos)

Uso (cd backend/):
    python scripts/explorar_docx.py                    # seções linkadas
    python scripts/explorar_docx.py --deep             # inclui sub-páginas
    python scripts/explorar_docx.py --media            # media library completa (+ órfãos)
    python scripts/explorar_docx.py --media --extrair  # baixa e extrai texto de tudo
    python scripts/explorar_docx.py --extrair          # baixa e extrai das seções linkadas
    python scripts/explorar_docx.py --url <url>        # extrai texto de um .docx específico
    python scripts/explorar_docx.py --busca wilczek    # filtra por nome/termo

Dependências extras além do requirements.txt do backend:
    pip install python-docx
"""
import argparse
import asyncio
import html as hlib
import io
import re
import sys
from pathlib import Path
from urllib.parse import urljoin

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

try:
    import httpx
except ImportError:
    sys.exit("Instale httpx: pip install httpx")

WP_BASE = "https://transparencia.caupr.gov.br/wp-json/wp/v2"

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Referer":         "https://transparencia.caupr.gov.br/",
    "Accept":          "*/*",
    "Accept-Language": "pt-BR,pt;q=0.9",
}

# Seções principais (mesmo mapeamento do scrape_transparencia_local.py)
SECOES_IDS: list[tuple[str, int]] = [
    ("Dispensa Eletrônica",           19317),
    ("Contratações Diretas",          19319),
    ("Atas de Registro de Preço",     19321),
    ("Relação de Contratos",          17867),
    ("Relação de Convênios",          17872),
    ("Relatórios ao TCU",             160),
    ("Relatórios e Pareceres",        17203),
    ("Auditoria Independente",        18731),
    ("Licitações / Editais",          341),
    ("Súmulas Conselho Diretor",      18692),
    ("Súmulas Comissões Temporárias", 18695),
    ("Resoluções",                    110),
    ("Atos Declaratórios",            118),
    ("Atos Declaratórios (antigos)",  17932),
    ("Orientações Jurídicas",         120),
    ("Pautas das Reuniões Plenárias", 18664),
    ("Folhas de Pagamento",           17580),
    ("Tabelas de Remuneração",        17225),
    ("Acordos Nacionais",             122),
    ("Acordos Internacionais",        124),
    ("Eleições do CAU",               597),
]

DOCX_RE = re.compile(r'\.docx?(\?[^"\']*)?$', re.IGNORECASE)


# ── Extração de links ─────────────────────────────────────────────────────────

def _extrair_links_docx(rendered: str, base_url: str = "") -> list[tuple[str, str]]:
    vistos: set[str] = set()
    out: list[tuple[str, str]] = []

    # Formato JSON do WP (aspas escapadas)
    for m in re.finditer(r'href=\\"([^\\"]+)\\"[^>]*>([^<]*)<', rendered):
        href = m.group(1).replace('\\/', '/')
        texto = hlib.unescape(m.group(2)).strip()
        if DOCX_RE.search(href) and href not in vistos:
            out.append((texto or _filename(href), href))
            vistos.add(href)

    # HTML normal
    for m in re.finditer(r'href="([^"]+)"[^>]*>([^<]*)<', rendered):
        href = m.group(1)
        if base_url and not href.startswith("http"):
            href = urljoin(base_url, href)
        texto = hlib.unescape(m.group(2)).strip()
        if DOCX_RE.search(href) and href not in vistos:
            out.append((texto or _filename(href), href))
            vistos.add(href)

    return out


def _filename(href: str) -> str:
    return href.split('?')[0].rstrip('/').split('/')[-1]


# ── Consulta WP REST API ──────────────────────────────────────────────────────

async def _fetch_page(http: httpx.AsyncClient, page_id: int) -> tuple[str, str]:
    """Retorna (rendered_content, page_link)."""
    url = f"{WP_BASE}/pages/{page_id}?_fields=content,link"
    try:
        r = await http.get(url, timeout=30)
        r.raise_for_status()
        data = r.json()
        return data.get("content", {}).get("rendered", ""), data.get("link", "")
    except Exception as e:
        print(f"    ⚠  API falhou (id={page_id}): {e}", file=sys.stderr)
        return "", ""


async def _fetch_child_pages(http: httpx.AsyncClient, parent_id: int) -> list[tuple[int, str]]:
    """Retorna [(child_id, child_title)] de todas as sub-páginas de parent_id."""
    url = f"{WP_BASE}/pages?parent={parent_id}&per_page=100&_fields=id,title,link"
    try:
        r = await http.get(url, timeout=30)
        r.raise_for_status()
        return [
            (p["id"], hlib.unescape(p.get("title", {}).get("rendered", f"ID {p['id']}")))
            for p in r.json()
        ]
    except Exception as e:
        print(f"    ⚠  Erro ao buscar sub-páginas de {parent_id}: {e}", file=sys.stderr)
        return []


# ── Extração de texto DOCX ────────────────────────────────────────────────────

def _extrair_texto_docx(docx_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "[python-docx não instalado — execute: pip install python-docx]"

    doc = Document(io.BytesIO(docx_bytes))
    partes: list[str] = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            partes.append(t)

    for table in doc.tables:
        for row in table.rows:
            celulas = [c.text.strip() for c in row.cells if c.text.strip()]
            if celulas:
                partes.append(" | ".join(celulas))

    return "\n".join(partes)


async def _baixar_e_extrair(http: httpx.AsyncClient, href: str) -> str:
    try:
        resp = await http.get(href, timeout=60)
        resp.raise_for_status()
        return _extrair_texto_docx(resp.content)
    except Exception as e:
        return f"[erro no download: {e}]"


# ── Descoberta principal ──────────────────────────────────────────────────────

async def descobrir(
    http:    httpx.AsyncClient,
    deep:    bool,
    busca:   str | None,
) -> list[dict]:
    """
    Retorna lista de {secao, titulo_link, href, filename}.
    Se deep=True, também varre sub-páginas de cada seção.
    """
    resultados: list[dict] = []
    urls_vistos: dict[str, str] = {}  # href → secao onde apareceu primeiro

    async def _processar(secao_label: str, page_id: int) -> None:
        rendered, page_link = await _fetch_page(http, page_id)
        links = _extrair_links_docx(rendered, page_link)
        for titulo, href in links:
            fn = _filename(href)
            item = {"secao": secao_label, "titulo": titulo, "href": href, "filename": fn}
            if href in urls_vistos:
                item["duplicado_de"] = urls_vistos[href]
            else:
                urls_vistos[href] = secao_label
            resultados.append(item)

        if deep:
            filhos = await _fetch_child_pages(http, page_id)
            if filhos:
                print(f"    ↳ {len(filhos)} sub-páginas encontradas", flush=True)
            for child_id, child_title in filhos:
                child_label = f"{secao_label} › {child_title}"
                rendered2, link2 = await _fetch_page(http, child_id)
                links2 = _extrair_links_docx(rendered2, link2)
                for titulo2, href2 in links2:
                    fn2 = _filename(href2)
                    item2 = {"secao": child_label, "titulo": titulo2, "href": href2, "filename": fn2}
                    if href2 in urls_vistos:
                        item2["duplicado_de"] = urls_vistos[href2]
                    else:
                        urls_vistos[href2] = child_label
                    resultados.append(item2)
                await asyncio.sleep(0.5)

    for i, (label, pid) in enumerate(SECOES_IDS, 1):
        print(f"  [{i:2d}/{len(SECOES_IDS)}] {label}...", end=" ", flush=True)
        await _processar(label, pid)
        n = sum(1 for r in resultados if r["secao"].startswith(label))
        print(f"{n} .docx" if n else "nenhum", flush=True)
        await asyncio.sleep(0.8)

    if busca:
        termo = busca.lower()
        return [r for r in resultados if termo in r["href"].lower() or termo in r["titulo"].lower()]

    return resultados


# ── Media Library (todos os uploads, incluindo órfãos) ───────────────────────

async def descobrir_media_library(
    http:   httpx.AsyncClient,
    busca:  str | None,
    extrair: bool,
) -> None:
    """
    Consulta /wp/v2/media?mime_type=application/vnd.openxmlformats-officedocument...
    para listar TODOS os .docx enviados ao WordPress, mesmo os não linkados.
    """
    MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    MIME_DOC  = "application/msword"

    resultados: list[dict] = []

    for mime in (MIME_DOCX, MIME_DOC):
        page = 1
        while True:
            url = f"{WP_BASE}/media?mime_type={mime}&per_page=100&page={page}&_fields=id,date,slug,source_url,title"
            try:
                r = await http.get(url, timeout=30)
                if r.status_code == 400:
                    break
                r.raise_for_status()
                items = r.json()
                if not items:
                    break
                for item in items:
                    href    = item.get("source_url", "")
                    titulo  = hlib.unescape(item.get("title", {}).get("rendered", ""))
                    data    = item.get("date", "")[:10]
                    fn      = _filename(href)
                    resultados.append({"titulo": titulo, "href": href, "filename": fn, "data": data})
                page += 1
                await asyncio.sleep(0.5)
            except Exception as e:
                print(f"  ⚠  Erro na página {page}: {e}", file=sys.stderr)
                break

    if busca:
        termo = busca.lower()
        resultados = [r for r in resultados if termo in r["href"].lower() or termo in r["titulo"].lower() or termo in r["filename"].lower()]

    print(f"\n{'='*65}")
    print(f"MEDIA LIBRARY: {len(resultados)} arquivos .doc/.docx encontrados")
    print(f"{'='*65}\n")

    for r in sorted(resultados, key=lambda x: x["data"]):
        print(f"  [{r['data']}]  {r['filename'][:65]}")
        print(f"            {r['href']}")
        if r["titulo"] and r["titulo"] != r["filename"]:
            print(f"            título: {r['titulo'][:60]}")

    if extrair and resultados:
        print(f"\n{'='*65}")
        print(f"EXTRAÇÃO DE TEXTO ({len(resultados)} arquivos)")
        print(f"{'='*65}\n")
        for i, r in enumerate(sorted(resultados, key=lambda x: x["data"]), 1):
            print(f"[{i}/{len(resultados)}] {r['filename']} ({r['data']})")
            texto = await _baixar_e_extrair(http, r["href"])
            if texto.strip():
                preview = texto[:400].replace('\n', ' ')
                print(f"  {len(texto):,} chars — {preview}...")
            else:
                print("  [sem texto extraível]")
            print()
            await asyncio.sleep(0.8)


# ── Extração de URL única ─────────────────────────────────────────────────────

async def extrair_url_unica(url: str) -> None:
    print(f"\nBaixando: {url}\n")
    async with httpx.AsyncClient(headers=HEADERS, follow_redirects=True) as http:
        texto = await _baixar_e_extrair(http, url)
    if not texto.strip():
        print("[documento vazio ou sem texto extraível]")
    else:
        print(texto)
        print(f"\n── Caracteres extraídos: {len(texto):,}")


# ── Main ──────────────────────────────────────────────────────────────────────

async def main(args: argparse.Namespace) -> None:
    if args.url:
        await extrair_url_unica(args.url)
        return

    async with httpx.AsyncClient(headers=HEADERS, follow_redirects=True) as http:
        if args.media:
            print(f"\n{'='*65}")
            print("Explorador .docx — Media Library completa (inclui órfãos)")
            print(f"{'='*65}")
            if args.busca:
                print(f"Filtro: '{args.busca}'")
            print()
            await descobrir_media_library(http, args.busca or None, args.extrair)
            return

    print(f"\n{'='*65}")
    print("Explorador .docx — Portal da Transparência CAU/PR")
    print(f"{'='*65}")
    print(f"Modo: {'deep (inclui sub-páginas)' if args.deep else 'normal'}")
    if args.busca:
        print(f"Filtro: '{args.busca}'")
    print()

    async with httpx.AsyncClient(headers=HEADERS, follow_redirects=True) as http:
        resultados = await descobrir(http, args.deep, args.busca)

        unicos    = [r for r in resultados if "duplicado_de" not in r]
        dupes     = [r for r in resultados if "duplicado_de" in r]
        unicas_urls = {r["href"] for r in unicos}

        print(f"\n{'='*65}")
        print(f"RESULTADO: {len(resultados)} ocorrências — {len(unicas_urls)} URLs únicas — {len(dupes)} duplicatas")
        print(f"{'='*65}\n")

        secao_atual = None
        for r in resultados:
            if r["secao"] != secao_atual:
                secao_atual = r["secao"]
                print(f"\n── {secao_atual}")

            dup_tag = f"  [DUP de: {r['duplicado_de']}]" if "duplicado_de" in r else ""
            print(f"   {r['filename'][:70]}{dup_tag}")
            print(f"   {r['href']}")

        # Duplicatas por nome de arquivo (arquivos diferentes, mesmo nome)
        from collections import Counter
        fn_count = Counter(r["filename"] for r in resultados)
        mesmo_nome = {fn: c for fn, c in fn_count.items() if c > 1}
        if mesmo_nome:
            print(f"\n── Arquivos com mesmo nome em seções diferentes:")
            for fn, c in sorted(mesmo_nome.items(), key=lambda x: -x[1]):
                print(f"   {fn}  ({c}x)")

        if args.extrair and unicas_urls:
            print(f"\n{'='*65}")
            print(f"EXTRAÇÃO DE TEXTO ({len(unicas_urls)} arquivos únicos)")
            print(f"{'='*65}\n")
            for i, r in enumerate([x for x in resultados if "duplicado_de" not in x], 1):
                print(f"[{i}/{len(unicas_urls)}] {r['filename']}")
                texto = await _baixar_e_extrair(http, r["href"])
                if texto.strip():
                    preview = texto[:500].replace('\n', ' ')
                    print(f"  {len(texto):,} chars — preview: {preview}...")
                else:
                    print("  [sem texto extraível]")
                print()
                await asyncio.sleep(1.0)


# ── Busca em conteúdo (grep interno dos docx) ─────────────────────────────────

async def busca_conteudo(
    http:       httpx.AsyncClient,
    termo:      str,
    ano_inicio: int,
    ano_fim:    int,
) -> None:
    """
    Baixa e extrai texto de todos os .docx da media library no intervalo de anos
    e procura 'termo' (case-insensitive) no conteúdo de cada arquivo.
    """
    MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    MIME_DOC  = "application/msword"

    todos: list[dict] = []
    for mime in (MIME_DOCX, MIME_DOC):
        page = 1
        while True:
            url = (
                f"{WP_BASE}/media?mime_type={mime}&per_page=100&page={page}"
                f"&after={ano_inicio}-01-01T00:00:00&before={ano_fim}-12-31T23:59:59"
                f"&_fields=id,date,source_url,title"
            )
            try:
                r = await http.get(url, timeout=30)
                if r.status_code == 400:
                    break
                r.raise_for_status()
                items = r.json()
                if not items:
                    break
                todos.extend(items)
                page += 1
                await asyncio.sleep(0.3)
            except Exception as e:
                print(f"  ⚠  Erro: {e}", file=sys.stderr)
                break

    print(f"\nTotal de arquivos em {ano_inicio}–{ano_fim}: {len(todos)}")
    print(f"Buscando '{termo}' no conteúdo de cada arquivo...\n")

    encontrados: list[dict] = []
    for i, item in enumerate(todos, 1):
        href   = item.get("source_url", "")
        data   = item.get("date", "")[:10]
        fn     = _filename(href)
        print(f"  [{i:3d}/{len(todos)}] {fn[:55]}...".ljust(70), end="\r", flush=True)
        texto = await _baixar_e_extrair(http, href)
        if termo.lower() in texto.lower():
            idx = texto.lower().find(termo.lower())
            ctx = texto[max(0, idx-100):idx+200].replace('\n', ' ')
            encontrados.append({"data": data, "filename": fn, "href": href, "ctx": ctx})
        await asyncio.sleep(0.5)

    print()
    print(f"\n{'='*65}")
    print(f"ENCONTRADO em {len(encontrados)}/{len(todos)} arquivos")
    print(f"{'='*65}\n")
    for r in encontrados:
        print(f"  [{r['data']}]  {r['filename']}")
        print(f"  URL: {r['href']}")
        print(f"  ...{r['ctx']}...")
        print()


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Explorador de .docx no Portal da Transparência do CAU/PR")
    parser.add_argument("--deep",    action="store_true", help="Varre sub-páginas de cada seção (ETPs, relatórios aninhados, etc.)")
    parser.add_argument("--media",   action="store_true", help="Consulta media library do WP — encontra TODOS os .docx, inclusive órfãos não linkados")
    parser.add_argument("--extrair", action="store_true", help="Baixa e extrai texto de cada .docx encontrado")
    parser.add_argument("--busca",   type=str, default="", help="Filtra resultados por nome/URL (metadata)")
    parser.add_argument("--grep",    type=str, default="", help="Busca DENTRO do conteúdo dos .docx (requer --ano)")
    parser.add_argument("--ano",     type=str, default="", help="Intervalo de anos para --grep, ex: 2019-2023 ou 2021")
    parser.add_argument("--url",     type=str, default="", help="Extrai texto de um .docx específico pelo URL")
    args = parser.parse_args()

    async def _main() -> None:
        if args.url:
            await extrair_url_unica(args.url)
            return

        async with httpx.AsyncClient(headers=HEADERS, follow_redirects=True) as http:
            if args.grep:
                ano_ini, ano_fim = 2019, 2026
                if args.ano:
                    partes = args.ano.split("-")
                    ano_ini = int(partes[0])
                    ano_fim = int(partes[-1])
                print(f"\n{'='*65}")
                print(f"Busca em conteúdo: '{args.grep}' ({ano_ini}–{ano_fim})")
                print(f"{'='*65}")
                await busca_conteudo(http, args.grep, ano_ini, ano_fim)
                return

            if args.media:
                print(f"\n{'='*65}")
                print("Explorador .docx — Media Library completa (inclui órfãos)")
                print(f"{'='*65}")
                if args.busca:
                    print(f"Filtro: '{args.busca}'")
                print()
                await descobrir_media_library(http, args.busca or None, args.extrair)
                return

            print(f"\n{'='*65}")
            print("Explorador .docx — Portal da Transparência CAU/PR")
            print(f"{'='*65}")
            print(f"Modo: {'deep (inclui sub-páginas)' if args.deep else 'normal'}")
            if args.busca:
                print(f"Filtro: '{args.busca}'")
            print()

            resultados = await descobrir(http, args.deep, args.busca)

            unicos      = [r for r in resultados if "duplicado_de" not in r]
            dupes       = [r for r in resultados if "duplicado_de" in r]
            unicas_urls = {r["href"] for r in unicos}

            print(f"\n{'='*65}")
            print(f"RESULTADO: {len(resultados)} ocorrências — {len(unicas_urls)} URLs únicas — {len(dupes)} duplicatas")
            print(f"{'='*65}\n")

            secao_atual = None
            for r in resultados:
                if r["secao"] != secao_atual:
                    secao_atual = r["secao"]
                    print(f"\n── {secao_atual}")
                dup_tag = f"  [DUP de: {r['duplicado_de']}]" if "duplicado_de" in r else ""
                print(f"   {r['filename'][:70]}{dup_tag}")
                print(f"   {r['href']}")

            from collections import Counter
            fn_count = Counter(r["filename"] for r in resultados)
            mesmo_nome = {fn: c for fn, c in fn_count.items() if c > 1}
            if mesmo_nome:
                print(f"\n── Arquivos com mesmo nome em seções diferentes:")
                for fn, c in sorted(mesmo_nome.items(), key=lambda x: -x[1]):
                    print(f"   {fn}  ({c}x)")

            if args.extrair and unicas_urls:
                print(f"\n{'='*65}")
                print(f"EXTRAÇÃO DE TEXTO ({len(unicas_urls)} arquivos únicos)")
                print(f"{'='*65}\n")
                for i, r in enumerate([x for x in resultados if "duplicado_de" not in x], 1):
                    print(f"[{i}/{len(unicas_urls)}] {r['filename']}")
                    texto = await _baixar_e_extrair(http, r["href"])
                    if texto.strip():
                        preview = texto[:500].replace('\n', ' ')
                        print(f"  {len(texto):,} chars — preview: {preview}...")
                    else:
                        print("  [sem texto extraível]")
                    print()
                    await asyncio.sleep(1.0)

    asyncio.run(_main())
